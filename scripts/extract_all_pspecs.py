"""
Author : Rajareddy Eddulannagari
Date: 12_Mar_2021

Description:
    This tool converts SRD docx to PSP file format for notepad++ (PSP : pseudocode file extension)
    this ensures hierarchy of pspecs, procedures and diagrams as per the SRD docx
	PSP files can be opened in Notepad++ with pseudocode language plugin (PSP) for pseudocode syntax highlighting"

Usage:
    Inputs:
        Rs_T_A400392-xx_accepted is SRD with one or all CRs markup selected, and used "Accept All changes shown" and saved
        Rs_T_A400392-xx_rejected is SRD with one or all CRs markup selected, and used "Reject All changes shown" and saved

    Commands:
        $python extract_all_pspecs.py Rs_T_A400392-xx_accepted.docx
        OR
        $python extract_all_pspecs.py Rs_T_A400392-xx_accepted.docx Rs_T_A400392-xx_rejected.docx

    Output:
        This tool will create "accepted" and / or "rejected" folders with all PSP files , based on number of docx files
        passed as arguments , as mentioned in Commands section above.
        Now these folders can be directly compared with beyond compare to see the impacted changes(select compare by content)
        You can add the pseudo code grammar for beyoncompare for better readabilty.

Python-docx info
    Each line in docx file is considered as paragraph in python-docx, this paragraph contains
    all styles related to that line of text.

"""

import inspect
import os
import docx
import shutil
import sys
import time
import datetime

heading_styles = ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6", "Heading 7",
                  "Heading 8", "Heading 9", "Title1", "Title2", "Title3"]
comment_styles = ["q", "a_b_c", "annotation text", "Caption", "dde", "der_el_a_b_c", "Footer", "gfd_coverage",
                  "Labelling", "List Paragraph", "No Spacing", "Normal_a_moi", "numero_dde", "Par_1er", "Par_espacé",
                  "Par_normal", "Par_sans_numéro", "Req_ID", "Req_text", "Revisions / Contents", "Table",
                  "table of figures", "Texte", "Title3", "toc 1", "toc 2", "toc 3", "toc 4", "toc 5", "toc 6", "toc 7",
                  "toc 8", "toc 9"]
pseudo_code_styles = ["pseudocode", "pseudocode Car Car Car",
                      "pseudocode Car Car Car Car", "Normal"]

open_loop_keywords = ["IF", "CASE", "FOR",
                      "WHILE", "ELSE", "ELSIF", "WHEN", "("]
close_loop_keyword = ["ENDIF", "ENDCASE", "ENDFOR",
                      "ENDWHILE", "ELSE", "ELSIF", "WHEN", ")"]

file_name_bad_chars = ['/', '\\', '"', ';', ':', '!', "*"]
hierarchy_diagrams_list = []


def update_diagram_hierarchy_list(current_diagram_number, current_diagram_name):
    no_of_digits_curr_diagram = sum(c.isdigit()
                                    for c in current_diagram_number)

    if no_of_digits_curr_diagram > 0:
        if len(hierarchy_diagrams_list) >= no_of_digits_curr_diagram:
            hierarchy_diagrams_list[no_of_digits_curr_diagram - 1] = (
                current_diagram_number + ' ' + current_diagram_name)

        else:
            hierarchy_diagrams_list.insert(no_of_digits_curr_diagram - 1,
                                           (current_diagram_number + ' ' + current_diagram_name))

        # hierarchy list cleanup of previous length
        while (no_of_digits_curr_diagram < len(hierarchy_diagrams_list)):
            index = (len(hierarchy_diagrams_list) - no_of_digits_curr_diagram)
            hierarchy_diagrams_list.pop(len(hierarchy_diagrams_list) - index)


def update_diagram_hierarchy_list_for_pspec(current_pspec_number, current_diagram_number):
    if (len(hierarchy_diagrams_list) > 1):
        no_of_digits_curr_pspec = sum(c.isdigit()
                                      for c in current_pspec_number)
        no_of_digits_curr_diagram = sum(
            c.isdigit() for c in hierarchy_diagrams_list[len(hierarchy_diagrams_list) - 1])
        # hierarchy list cleanup of previous length
        if (no_of_digits_curr_diagram >= no_of_digits_curr_pspec):
            no_of_pops = (no_of_digits_curr_diagram -
                          no_of_digits_curr_pspec) + 1
            while (no_of_pops):
                no_of_pops = no_of_pops - 1
                if (len(hierarchy_diagrams_list) > 1):
                    hierarchy_diagrams_list.pop(
                        len(hierarchy_diagrams_list) - 1)


# This procedure extracts all pspec's under matching diagram .
def extract_diagrams_and_pspecs(docx_file, dir_name):
    current_pspec_number = r'P-Spec x.y..'
    current_pspec_name = r'name_not_found'
    current_diagram_number = r'Diagram x.y..'
    current_diagram_name = r'name_not_found'
    open_loop_count = 0
    space_str = ""
    alignement = "    "
    is_it_procedure = False

    doc = docx.Document(docx_file)
    print("SRD docx is being parsed : It may take around 2-mins")

    for para in doc.paragraphs:
        if para.style.name in heading_styles:
            str_list = para.text.split(';')
            if is_diagram(str_list[0].strip()):
                # holds heading before ";" string
                current_diagram_number = str_list[0].strip()
                current_pspec_number = current_diagram_number  # reset on every diagram detection
                if (len(str_list) > 1):
                    current_diagram_name = str_list[1].strip()
                    for ch in file_name_bad_chars:
                        current_diagram_name = current_diagram_name.replace(
                            ch, '')
                    current_pspec_name = 'desciption'  # reset on every diagram detection
                # print("\t\t\t", current_heading_diagram, diagram)
                update_diagram_hierarchy_list(
                    current_diagram_number, current_diagram_name)
                is_it_procedure = False
                open_loop_count = 0

            if is_pspec(str_list[0].strip()):
                # holds heading before ";" string
                current_pspec_number = str_list[0].strip()
                if (len(str_list) > 1):
                    # holds heading after ";" string
                    current_pspec_name = str_list[1].strip()
                    for ch in file_name_bad_chars:
                        current_pspec_name = current_pspec_name.replace(ch, '')
                # print("\t\t",current_heading_pspec.strip(),"::",current_heading_name.strip())
                update_diagram_hierarchy_list_for_pspec(
                    current_pspec_number, current_diagram_number)
                is_it_procedure = False
                open_loop_count = 0

            if is_procedure(str_list[0].strip()):
                is_it_procedure = True
                open_loop_count = 0

        # update dir_path as per current hierarchy
        dir_path = dir_name
        for name in hierarchy_diagrams_list:
            dir_path = os.path.join(dir_path, name)
            # dir_path = dir_path.replace('Diagram', 'D')
            # dir_path = dir_path.replace(' ', '_')

            if not (os.path.exists(dir_path)):
                os.mkdir(dir_path)

        # write_file_name = (current_pspec_number + current_pspec_name).replace('P-Spec', 'P')
        write_file_name = (current_pspec_number + ' ' + current_pspec_name)
        # write_file_name = write_file_name.replace(' ', '_') + ".pseudo"
        write_file_name = write_file_name + ".pseudo"

        if is_it_procedure:
            # procedures are at root folder level
            write_file_path = os.path.join(dir_name, write_file_name)
        else:
            write_file_path = os.path.join(dir_path, write_file_name)

        with open(write_file_path, "a+", encoding="utf-8") as out_file:
            if len(para.text) > 0:
                if (para.style.name in pseudo_code_styles):  # pseudo code syntax
                    temp = para.text.split(" ")
                    if (len(temp) > 0):
                        if (temp[0].strip() in close_loop_keyword) and (open_loop_count > 0):
                            open_loop_count = open_loop_count - 1
                        space_str = "{}".format(alignement * open_loop_count)
                        if (temp[0].strip() in open_loop_keywords):
                            open_loop_count = open_loop_count + 1
                    # in IF statement & conditions were not aligned correctly hence replaced the newline
                    line = "{}{}".format(
                        space_str, para.text.replace("\n", "\n" + space_str))
                elif current_pspec_number in para.text:
                    line = "{}{}{}".format('', "", para.text.replace(
                        "\n", " "))  # remove newline in comments
                else:
                    # remove newline in comments
                    line = "{}{}{}".format(
                        space_str, "// ", para.text.replace("\n", " "))
                print(line, file=out_file)


def is_diagram(first_str):
    if type(first_str) == str:
        str_diagram_space = first_str.split(" ")
        str_diagram_dot = first_str.split(".")
        if ("Diagram".lower() == str_diagram_space[0].lower()) or ("Diagram".lower() == str_diagram_dot[0].lower()):
            return True
        else:
            return False
    else:
        return False


def is_pspec(first_str):
    if type(first_str) == str:
        str_diagram_space = first_str.split(" ")
        str_diagram_dot = first_str.split(".")
        if ("P-Spec".lower() == str_diagram_space[0].lower()) or (
                "Procedure".lower() == str_diagram_space[0].lower()) or \
                ("P-Spec".lower() == str_diagram_dot[0].lower()) or ("Procedure".lower() == str_diagram_dot[0].lower()):
            return True
        else:
            return False
    else:
        return False


def is_procedure(first_str):
    if type(first_str) == str:
        str_diagram_space = first_str.split(" ")
        str_diagram_dot = first_str.split(".")
        if ("Procedure".lower() == str_diagram_space[0].lower()) or ("Procedure".lower() == str_diagram_dot[0].lower()):
            return True
        else:
            return False
    else:
        return False


def delete_dir_and_add_dir(script_dir, dir):
    path = os.path.join(script_dir, dir)
    print(path)
    if os.path.exists(path):
        shutil.rmtree(path)
        print("% s has been removed successfully" % dir)
    time.sleep(1)
    if not os.path.exists(path):
        os.mkdir(path)
        print("% s has been created successfully" % dir)


def check_if_current_diagram_in_specified_list(diagram_number, pspec_name_list):
    to_remove_chars = ['.', " "]
    if type(diagram_number) == str:
        for ch in to_remove_chars:
            diagram_number = diagram_number.replace(ch, '')
        diagram_number = diagram_number.lower()
        # print("diagram_number : ",diagram_number )
        if (diagram_number.strip() in pspec_name_list):
            return True
        else:
            return False
    else:
        return False


def check_if_current_pspec_in_specified_list(pspec_number, pspec_name_list):
    to_remove_chars = ['.', " "]
    if type(pspec_number) == str:
        for ch in to_remove_chars:
            pspec_number = pspec_number.replace(ch, '')
        pspec_number = pspec_number.lower()
        # print( "pspec_number : ",pspec_number)
        if (pspec_number.strip() in pspec_name_list):
            return True
        else:
            return False
    else:
        return False


# This function is for development / debug.  which list all style names in a documents
def list_all_styles_in_docx(docx_file):
    doc = docx.Document(docx_file)
    with open("style_name.txt", "a+", encoding="utf-8") as out_file:
        for para in doc.paragraphs:
            if len(para.style.name) > 0:
                print(para.style.name, file=out_file)


def main(argv):
    # Get the SRD names , where changes are accepted in one and changes are rejected in another for a CR
    no_of_arguments = len(argv)

    script_dir = os.path.dirname(os.path.abspath(
        inspect.getfile(inspect.currentframe())))

    # Extracting the specified pspecs from SRD with accepted changes for a CR
    ct = datetime.datetime.now()
    print("current time:-", ct)

    srd_file_name_accepted = argv[0].strip()
    delete_dir_and_add_dir(script_dir, "accepted")
    srd_file_path = os.path.join(script_dir, srd_file_name_accepted)
    dir_name = os.path.join(script_dir, "accepted")
    extract_diagrams_and_pspecs(srd_file_path, dir_name)

    ct = datetime.datetime.now()
    print("current time:-", ct)

    if (no_of_arguments > 1):
        srd_file_name_rejected = argv[1].strip()
        # Extracting the specified pspecs from SRD with rejected changes for a CR
        delete_dir_and_add_dir(script_dir, "rejected")
        srd_file_path = os.path.join(script_dir, srd_file_name_rejected)
        dir_name = os.path.join(script_dir, "rejected")
        extract_diagrams_and_pspecs(srd_file_path, dir_name)

        ct = datetime.datetime.now()
        print("current time:-", ct)


# End of main

if __name__ == "__main__":
    if len(sys.argv) == 2 or len(sys.argv) == 3:
        main(sys.argv[1:])
    else:
        print("usage is : \n"
              "python extract_all_pspecs.py Rs_T_Axxx392-xx_accepted.docx \n" "OR \n"
              "python extract_all_pspecs.py Rs_T_Axxx392-xx_accepted.docx Rs_T_Axxx392-xx_rejected.docx")
